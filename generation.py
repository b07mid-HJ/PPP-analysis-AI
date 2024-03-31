import streamlit as st
from streamlit_sortables import sort_items
import langchain_doc_gen as lg
from io import BytesIO
from docx.oxml.shared import qn  # feel free to move these out
from docx.oxml import OxmlElement
from docx.shared import RGBColor
"""
# File Uploader
Upload a excel file here to generate a document with the sorted dataset.
"""

w = st.file_uploader("Upload a excel file", type="xlsx")

if w:
    from eparse.core import get_df_from_file
    from openpyxl import load_workbook
    import numpy as np

    wb = load_workbook(w)
    sheet_names = wb.sheetnames

    dataset = dict()
    tables = [table for table in get_df_from_file(w , sheet = sheet_names)]
        
    for tab in tables:
        tab_value = tab[0]
        tab_sheet = tab[3]
        arr = np.array(tab_value)
        tab_title = arr[-1, 0]
        row_table = {tab_title: tab_value}
        
        if tab_sheet not in dataset:
            dataset[tab_sheet] = [row_table]
        else :
            dataset[tab_sheet].append(row_table)

    original_items = []
    for sheet, tables in dataset.items():
        items = [key for table in tables for key in table.keys()]
        original_items.append({'header': sheet, 'items': items})

    sorted_items = sort_items(original_items, multi_containers=True, direction='vertical')

    if st.button('Generate Document'):
        # Initialize sorted_dataset to store the re-ordered dataset
        sorted_dataset = {}

        # Loop through each sorted item (sheet) in sorted_items
        for sorted_sheet in sorted_items:
            sheet_name = sorted_sheet['header']
            sorted_tables_titles = sorted_sheet['items']
            
            # Retrieve the unsorted tables for this sheet from the original dataset
            unsorted_tables = dataset[sheet_name]
            
            # Create a temporary dictionary to easily access tables by their title
            title_to_table_dict = {list(table.keys())[0]: table for table in unsorted_tables}
            
            # Initialize a list to hold the sorted tables for this sheet
            sorted_tables_for_sheet = []
            
            # Loop through each sorted table title and append the corresponding table
            # from title_to_table_dict to the sorted list
            for title in sorted_tables_titles:
                sorted_table = title_to_table_dict[title]
                sorted_tables_for_sheet.append(sorted_table)
            
            # Update the sorted_dataset with the sorted tables for this sheet
            sorted_dataset[sheet_name] = sorted_tables_for_sheet
            import pandas as pd
            import numpy as np
            import docx

            # Function to convert DataFrame into a list of lists without headers and the last row, handling NaN values
            def dataframe_to_list_without_headers_and_last_row(data):
                # Exclude the last row from the dataframe
                data_without_last_row = data.iloc[:-1]
                # Convert the dataframe into a list of lists without including headers and handling NaN values
                data_rows = data_without_last_row.values.tolist()
                # Replace NaN values with an empty string in the data rows
                cleaned_data_rows = [[str(cell) if pd.notnull(cell) else '' for cell in row] for row in data_rows]
                return cleaned_data_rows

            # Create a new Document
            doc = docx.Document()

            # Iterate through the sorted_dataset to add data to the Word document
            for sheet, tables in sorted_dataset.items():
                doc.add_heading(sheet, level=1)  # Sheet name as a title
                desc_list=[]
                for table in tables:
                    for title, data in table.items():
                        doc.add_heading(title, level=2)  # Table name as a subtitle
                        
                        # Check if data is a pandas DataFrame and convert if necessary
                        if isinstance(data, pd.DataFrame):
                            data_list = dataframe_to_list_without_headers_and_last_row(data)
                        else:
                            # Assuming data is already a list of lists. Exclude the last row.
                            # Make sure to handle NaN values and other transformations as needed.
                            data_list = data[:-1] if data else data

                        # Add a table to the document, making sure data_list is not empty
                        if data_list:
                            word_table = doc.add_table(rows=len(data_list), cols=len(data_list[0]) if data_list else 0)
                            word_table.style = 'TableGrid'
                            
                            # Populate the table with data, excluding the last row
                            for i, row in enumerate(data_list):
                                cells = [c for c in row if c != '']
                                brk = len(cells) == 1
                                for j, cell in enumerate(row):
                                    word_table.cell(i, j).text = cell
                                    if i == 0 or brk :
                                        clr = "#D9D9D9" if brk else "#2F5496"
                                        cell_properties = word_table.cell(i, j)._element.tcPr
                                        try:
                                            cell_shading = cell_properties.xpath('w:shd')[0]  # in case there's already shading
                                        except IndexError:
                                            cell_shading = OxmlElement('w:shd') # add new w:shd element to it
                                            cell_shading.set(qn('w:fill'), clr)  # set fill property, respecting namespace
                                        cell_properties.append(cell_shading)  # finally extend cell props with shading element
                                        if i == 0:
                                            word_table.cell(i, j).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # Blue color
                        doc.add_paragraph()
                        desc=lg.table_desc(title,data_list)
                        desc_list.append({"title":title,"desc":desc})
                        doc.add_paragraph(desc)  # Add an empty paragraph after each table    
                        doc.add_paragraph()
                doc.add_heading("Conclusion", level=2)
                doc.add_paragraph(lg.doc_content(sheet,desc_list))    
            # Save the document
            doc_io = BytesIO()
            # Save the document to the BytesIO object
            doc.save(doc_io)
            # Seek to the start of the BytesIO object so it can be read from the beginning
            doc_io.seek(0)
        st.write("Document generated successfully 🎉")
        st.download_button(
            label="Download Document",
            data=doc_io,  # Pass the BytesIO object here
            file_name="sorted_dataset.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )